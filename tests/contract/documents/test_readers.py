"""Contract tests for PDF and DOCX extraction adapters."""

from __future__ import annotations

from io import BytesIO

import fitz
from docx import Document

from multi_agentic_graph_rag.adapters.documents.docx_reader import DocxReader
from multi_agentic_graph_rag.adapters.documents.pdf_reader import PdfReader


def test_pdf_reader_preserves_page_provenance() -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "PDF paragraph")

    blocks = PdfReader().read(document.tobytes())

    assert [block.text for block in blocks] == ["PDF paragraph"]
    assert blocks[0].page_number == 1
    assert blocks[0].block_index == 0
    document.close()


def test_docx_reader_preserves_paragraph_provenance() -> None:
    document = Document()
    document.add_paragraph("")
    document.add_paragraph("DOCX paragraph")
    buffer = BytesIO()
    document.save(buffer)

    blocks = DocxReader().read(buffer.getvalue())

    assert [block.text for block in blocks] == ["DOCX paragraph"]
    assert blocks[0].paragraph_index == 1
    assert blocks[0].block_index == 0
